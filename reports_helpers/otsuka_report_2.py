"""Build an Otsuka-style media-coverage report (.docx) from dashboard data.

The output mirrors the reference report (`otsuka_report_example.docx`): a Calibri
11pt document with no logo, a bold title line, an executive-summary intro and
sign-off, then one bold section heading per media-monitoring section followed by
its articles. Each article renders as:

    Source/Author: Article Title (MM.DD.YY) (UVM: 50.59M)   (title is a blue link)
      • The article mentions Otsuka but there was no mention of ... in this coverage.
      • <summary point 1>
      • <summary point 2>
      • Also covered by Endpoints News, STAT, and ClinicalTrials Arena.

The mention note, every summary point, and the "Also covered by" line are all
bullets of one list, matching the reference document.

Sections with no articles render a single "NA" line. See
`build_media_monitoring_report` for the entry point. `sections_from_charts_data`
(shared with the BeOne report) extracts the ordered sections from charts data.
"""

from __future__ import annotations

import io
from datetime import date, datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

# Re-export so callers can build sections the same way as the other reports.
from reports_helpers.beone_report import (
    sections_from_charts_data, title_case, clean_author,
)

# --- constants ---------------------------------------------------------------

FONT = "Calibri"
FONT_SZ = Pt(11)  # matches the reference document's "No Spacing" body text.
FOOTER_SZ = Pt(9)  # the reference footnotes are 9pt.

# Hyperlinks use Word's built-in "Hyperlink" look (blue, underlined).
LINK_BLUE = RGBColor(0x05, 0x63, 0xC1)

FOOTER_COPYRIGHT = (
    "Materials included or referenced in this report may be the subject of "
    "copyright. Users must verify that they have sufficient permissions, rights, "
    "and/or licenses prior to any copying or redistribution. Please consult your "
    "copyright compliance guidelines and personnel with questions."
)
FOOTER_TIP = "Tip: Having trouble viewing a webpage? Try using a different browser."

DEFAULT_SIGNOFF = ("Thanks,", "Infovision Team")


# --- low-level docx helpers --------------------------------------------------

def _style_run(run, *, bold=False, italic=False, color=None, size=FONT_SZ,
               underline=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.font.underline = underline
    if color is not None:
        run.font.color.rgb = color
    return run


def _add_run(paragraph, text, **kw):
    return _style_run(paragraph.add_run(text), **kw)


def _add_external_hyperlink(paragraph, url, text, *, bold=False, italic=False,
                            color=LINK_BLUE, underline=True):
    """Add a clickable external hyperlink run to `paragraph`."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = paragraph.add_run(text)
    _style_run(run, bold=bold, italic=italic, color=color, underline=underline)
    hyperlink.append(run._element)
    paragraph._p.append(hyperlink)
    return hyperlink


def _tight(paragraph, *, space_before=0, space_after=0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return paragraph


def _blank_line(doc):
    p = doc.add_paragraph()
    _tight(p)
    _add_run(p, "")
    return p


def _bullet_paragraph(doc, *, space_after=2):
    """A bulleted list paragraph (Word's "List Bullet"), like the reference's
    per-article summary points. Falls back to a plain paragraph if the style is
    missing from the template."""
    try:
        p = doc.add_paragraph(style="List Bullet")
    except KeyError:
        p = doc.add_paragraph()
    _tight(p, space_after=space_after)
    # Indent the whole bullet one tab (0.5") in from the article header.
    p.paragraph_format.left_indent = Inches(0.5)
    return p


# --- value formatting --------------------------------------------------------

def _parse_date(value):
    if value is None or value == "":
        return None
    if isinstance(value, (datetime, date)):
        return value
    s = str(value).strip()
    # Normalize an ISO 8601 'Z' / timezone suffix so fromisoformat accepts it.
    iso = s.replace("Z", "+00:00")
    try:
        return datetime.fromisoformat(iso)
    except ValueError:
        pass
    for fmt in ("%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d",
                "%Y/%m/%d", "%d-%m-%Y", "%m/%d/%Y",
                "%d-%m-%y", "%m-%d-%y",
                "%m.%d.%Y", "%m.%d.%y", "%d.%m.%Y", "%Y.%m.%d"):
        try:
            return datetime.strptime(s[:len(fmt) + 2].strip(), fmt)
        except ValueError:
            continue
    try:
        return datetime.strptime(s[:10], "%Y-%m-%d")
    except ValueError:
        return None


def _fmt_article_date(value):
    """Format an article date as the reference's `MM.DD.YY` parenthetical."""
    dt = _parse_date(value)
    return dt.strftime("%m.%d.%y") if dt else ""


def _fmt_reach(value):
    """Humanize a reach/UVM count, e.g. 50_590_000 -> `50.59M`, 289_000 -> `289K`."""
    try:
        n = int(float(value))
    except (TypeError, ValueError):
        return ""
    if n <= 0:
        return ""
    if n >= 1_000_000:
        out = f"{n / 1_000_000:.2f}".rstrip("0").rstrip(".")
        return f"{out}M"
    if n >= 1_000:
        out = f"{n / 1_000:.2f}".rstrip("0").rstrip(".")
        return f"{out}K"
    return str(n)


def _article_source(article):
    src = (
        article.get("domain")
        or article.get("source")
        or article.get("publication")
        or article.get("domain_name")
        or "Unknown source"
    )
    return str(src).strip()


def _join_keywords(words):
    """Comma-separate with an Oxford `and`: [a] -> 'a', [a,b] -> 'a and b',
    [a,b,c] -> 'a, b, and c'."""
    words = [w for w in words if w]
    if not words:
        return ""
    if len(words) == 1:
        return words[0]
    if len(words) == 2:
        return f"{words[0]} and {words[1]}"
    return f"{', '.join(words[:-1])}, and {words[-1]}"


def _mention_note(article, keywords):
    """The article's brand-mention line. Prefers an LLM-written ``mention_note``
    on the article; otherwise builds the 'mentions X but no mention of Y' line
    from the brand/competitor keywords found in the article's title + content.
    Returns None when neither is available."""
    llm_note = str(article.get("mention_note") or "").strip()
    if llm_note:
        return llm_note

    keywords = [str(k).strip() for k in (keywords or []) if str(k).strip()]
    if not keywords:
        return None
    # De-dupe case-insensitively while preserving order.
    seen, unique = set(), []
    for kw in keywords:
        low = kw.lower()
        if low not in seen:
            seen.add(low)
            unique.append(kw)

    haystack = " ".join(
        str(article.get(f) or "") for f in ("title", "content", "summary")
    ).lower()
    mentioned = [kw for kw in unique if kw.lower() in haystack]
    missing = [kw for kw in unique if kw.lower() not in haystack]

    if mentioned and missing:
        return (f"The article mentions {_join_keywords(mentioned)} but there was "
                f"no mention of {_join_keywords(missing)} in this coverage.")
    if mentioned:
        return f"The article mentions {_join_keywords(mentioned)} in this coverage."
    return f"There was no mention of {_join_keywords(unique)} in this coverage."


def _summary_paragraphs(article):
    """Split an article's summary/content into display paragraphs (blank lines or
    single newlines start a new paragraph), preserving the reference's bulleted feel."""
    content = article.get("summary") or article.get("content") or ""
    content = str(content).strip()
    if not content:
        return []
    parts = [p.strip() for p in content.replace("\r\n", "\n").split("\n")]
    return [p for p in parts if p]


# --- report builder ----------------------------------------------------------

def build_media_monitoring_report(
    sections,
    *,
    report_date=None,
    brand="Otsuka",
    title=None,
    intro=None,
    summary_paragraphs=None,
    signoff=DEFAULT_SIGNOFF,
    brand_keywords=None,
    competitor_keywords=None,
):
    """Render the Otsuka news-coverage report and return the .docx as bytes.

    Args:
        sections: list of ``{"name": str, "articles": [article dict, ...]}`` in
            display order. Each article dict may contain title, content/summary,
            url, domain/source, author, date, reach and similar_articles
            ({domain: url}).
        report_date: date shown in the title line. Defaults to the latest article
            date, then today.
        brand: brand name used in the default title and intro.
        title: override the bold title line.
        intro: override the default intro sentence.
        summary_paragraphs: optional list of executive-summary paragraphs shown
            between the intro and the "Please find the full list..." line.
        signoff: trailing sign-off lines (e.g. ("Thanks,", "Infovision Team")).
        brand_keywords / competitor_keywords: used to build each article's
            "mentions X but no mention of Y" note; omit to skip that line.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = FONT_SZ

    if report_date is None:
        report_date = datetime.today()

    keywords = list(brand_keywords or []) + list(competitor_keywords or [])
    date_label = _fmt_article_date(report_date)

    # 1) Title line — centered, bold, underlined.
    if title is None:
        title = f"Otsuka IRA Announcement News Coverage – IPAY 2028 Cohort"
        if date_label:
            title = f"{title} {date_label}"
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight(tp, space_after=4)
    _add_run(tp, title, bold=True, underline=True)
    _blank_line(doc)

    # 2) Intro line.
    if intro is None:
        intro = (
            f"Please find below a summary of recent news coverage relevant to "
            f"{brand}, its pipeline and treatments, and the broader policy and "
            "competitive landscape."
        )
    ip = doc.add_paragraph()
    _tight(ip)
    _add_run(ip, intro)
    _blank_line(doc)

    # 3) Optional executive-summary paragraphs.
    for para in (summary_paragraphs or []):
        if not str(para).strip():
            continue
        sp = doc.add_paragraph()
        _tight(sp)
        _add_run(sp, str(para).strip())
        _blank_line(doc)

    # 4) "Please find the full list of coverage below." + sign-off.
    fp = doc.add_paragraph()
    _tight(fp)
    _add_run(fp, "Please find the full list of coverage below.")
    _blank_line(doc)
    for line in (signoff or ()):
        lp = doc.add_paragraph()
        _tight(lp)
        _add_run(lp, str(line))
    _blank_line(doc)

    # 5) Sections.
    for section in sections:
        articles = section.get("articles", [])

        head = doc.add_paragraph()
        _tight(head, space_before=4, space_after=2)
        _add_run(head, str(section["name"]), bold=True, underline=True)

        if not articles:
            na = doc.add_paragraph()
            _tight(na, space_after=2)
            _add_run(na, "N/A")
            _blank_line(doc)
            continue

        for art in articles:
            # Header line — "Source/Author: Title (MM.DD.YY) (UVM: reach)".
            hp = doc.add_paragraph()
            _tight(hp, space_after=1)
            source = _article_source(art)
            author = clean_author(art.get("author")) or "NA"
            label = f"{source}/{author}: "
            _add_run(hp, label)

            title_text = title_case(art.get("title") or "Untitled")
            url = art.get("url") or ""
            if url:
                _add_external_hyperlink(hp, url, title_text)
            else:
                _add_run(hp, title_text)

            art_date = _fmt_article_date(art.get("date"))
            if art_date:
                _add_run(hp, f" ({art_date})")
            reach = _fmt_reach(art.get("reach"))
            if reach:
                _add_run(hp, f" (UVM: {reach})")

            # Brand-mention note — first bullet of the article's summary list.
            note = _mention_note(art, keywords)
            if note:
                _add_run(_bullet_paragraph(doc), note)

            # Summary points — one bullet per point.
            for para in _summary_paragraphs(art):
                _add_run(_bullet_paragraph(doc), para)

            # "Also covered by ..." — the article's last bullet, comma-separated links.
            similar = art.get("similar_articles") or {}
            similar = similar if isinstance(similar, dict) else {}
            if similar:
                ap = _bullet_paragraph(doc)
                _add_run(ap, "Also covered by ")
                items = list(similar.items())
                last = len(items) - 1
                for j, (domain, durl) in enumerate(items):
                    if j:
                        if len(items) > 2:
                            _add_run(ap, ", and " if j == last else ", ", italic=True)
                        else:
                            _add_run(ap, " and ", italic=True)
                    if durl:
                        _add_external_hyperlink(ap, durl, domain)
                    else:
                        _add_run(ap, domain)
                _add_run(ap, ".", italic=True)

            _blank_line(doc)

    # 6) Footer notes.
    cp = doc.add_paragraph()
    _tight(cp, space_before=10)
    _add_run(cp, FOOTER_COPYRIGHT, italic=True, size=FOOTER_SZ)
    tp = doc.add_paragraph()
    _tight(tp)
    _add_run(tp, FOOTER_TIP, italic=True, size=FOOTER_SZ)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()