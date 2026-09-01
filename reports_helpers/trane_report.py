"""Build a Trane-style media-monitoring "News Monitor" report (.docx).

The output mirrors the reference report (`trane_report_example.docx`): an Arial
9pt document with a banner logo, an intro line (with a mailto link), a report
date, then one purple ALL-CAPS section heading per media-monitoring section
followed by its articles. Each article renders as:

    SOURCE LABEL            (bold + italic, uppercase)
    Article title           (blue hyperlink when a URL is present, else plain)
    Summary:                (bold label)
    <summary text>
    This story was also covered by <domain>, <domain>.   (when similar copies)
    Translated from <language>.                          (for translated items)

See `build_media_monitoring_report` for the entry point. `sections_from_charts_data`
(shared with the BeOne report) extracts the ordered sections from charts data.
"""

from __future__ import annotations

import io
import os
from datetime import date, datetime

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Emu, Inches

# Re-export so callers can build sections the same way as the BeOne report.
from reports_helpers.beone_report import sections_from_charts_data, title_case

# --- constants ---------------------------------------------------------------

FONT = "Arial"
FONT_SZ = Pt(9)  # matches the reference document's Normal style.

# Section heading look (Word "Heading 1" in the reference): bold, all-caps,
# purple, 12pt.
HEADING_COLOR = RGBColor(0x64, 0x00, 0xFF)
HEADING_SZ = Pt(12)

# Subtle gray for the report date (the reference uses the "Subtle Reference"
# character style — small caps, muted color).
DATE_COLOR = RGBColor(0x5A, 0x5A, 0x5A)

# Hyperlinks use Word's "Hyperlink" look from the reference: bold, cyan,
# underlined. The "also covered by" links reuse the same style but force the
# color to black (color="auto" in the reference).
LINK_CYAN = RGBColor(0x00, 0xB9, 0xE4)
BLACK = RGBColor(0x00, 0x00, 0x00)

LOGO_PATH = os.path.join(os.path.dirname(__file__), "assets", "trane_logo.jpeg")
# Reference banner dimensions (EMU) — ~6.5in x 1.06in.
LOGO_WIDTH = Emu(5943600)
LOGO_HEIGHT = Emu(971550)

DEFAULT_SUBMIT_EMAIL = "communications@tranetechnologies.com"


# --- low-level docx helpers --------------------------------------------------

def _style_run(run, *, bold=False, italic=False, color=None, size=FONT_SZ,
               caps=False, small_caps=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if caps:
        rpr.append(OxmlElement("w:caps"))
    if small_caps:
        rpr.append(OxmlElement("w:smallCaps"))
    if color is not None:
        run.font.color.rgb = color
    return run


def _add_run(paragraph, text, **kw):
    return _style_run(paragraph.add_run(text), **kw)


def _add_external_hyperlink(paragraph, url, text, *, bold=True, italic=False,
                            color=LINK_CYAN, underline=True):
    """Add a clickable external hyperlink run to `paragraph`."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = paragraph.add_run(text)
    run.font.underline = underline
    _style_run(run, bold=bold, italic=italic, color=color)
    hyperlink.append(run._element)
    paragraph._p.append(hyperlink)
    return hyperlink


def _tight(paragraph, *, space_before=2, space_after=2):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return paragraph


def _blank_line(container):
    p = container.add_paragraph()
    _tight(p, space_before=0, space_after=0)
    _add_run(p, "")
    return p


def _no_table_borders(table):
    """Strip all table borders (the reference shows only Word's faint gridlines)."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


class _Block:
    """Writes paragraphs into one table-row cell, reusing the cell's initial
    empty paragraph for the first line so rows don't start with a blank gap."""

    def __init__(self, cell):
        self.cell = cell
        self._first_used = False

    def add_paragraph(self):
        if not self._first_used:
            self._first_used = True
            return self.cell.paragraphs[0]
        return self.cell.add_paragraph()


# --- value formatting --------------------------------------------------------

def _parse_date(value):
    if value is None or value == "":
        return None
    if isinstance(value, (datetime, date)):
        return value
    s = str(value)
    for fmt in ("%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d",
                "%Y/%m/%d", "%d-%m-%Y", "%m/%d/%Y"):
        try:
            return datetime.strptime(s[:len(fmt) + 2].strip(), fmt)
        except ValueError:
            continue
    try:
        return datetime.strptime(s[:10], "%Y-%m-%d")
    except ValueError:
        return None


def _fmt_report_date(value):
    """Format the report header date, e.g. `July 15, 2025`."""
    dt = _parse_date(value)
    if not dt:
        return ""
    return f"{dt.strftime('%B')} {dt.day}, {dt.year}"


def _article_source(article):
    src = (
        article.get("source")
        or article.get("publication")
        or article.get("domain_name")
        or article.get("domain")
        or "Unknown source"
    )
    return str(src).strip().upper()


def _translated_from(article):
    """The source language for a translated article, if any (else None)."""
    val = article.get("translated_from") or article.get("source_language")
    if val:
        return str(val).strip()
    lang = article.get("language")
    if lang and str(lang).strip().lower() not in ("", "en", "english", "en-us"):
        return str(lang).strip()
    return None


# --- report builder ----------------------------------------------------------

def build_media_monitoring_report(
    sections,
    *,
    report_date=None,
    brand="Trane Technologies",
    submit_email=DEFAULT_SUBMIT_EMAIL,
    intro=None,
    include_logo=True,
):
    """Render the Trane News Monitor report and return the .docx as bytes.

    Args:
        sections: list of ``{"name": str, "articles": [article dict, ...]}`` in
            display order. Each article dict may contain title, content/summary,
            url, domain/source/publication, date, similar_articles ({domain: url})
            and an optional translated_from/language.
        report_date: date shown in the header. Defaults to the latest article
            date, then today.
        brand: brand name used in the intro line and heading text.
        submit_email: contact email shown in the intro (linked as mailto:).
        intro: override the default intro sentence.
        include_logo: embed the banner logo if the asset is present.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = FONT_SZ

    if report_date is None:
        latest = None
        for section in sections:
            for art in section.get("articles", []):
                dt = _parse_date(art.get("date"))
                if dt and (latest is None or dt > latest):
                    latest = dt
        report_date = latest or datetime.today()

    # 1) Logo banner.
    if include_logo and os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        _tight(p, space_before=0, space_after=6)
        run = p.add_run()
        try:
            run.add_picture(LOGO_PATH, width=LOGO_WIDTH, height=LOGO_HEIGHT)
        except Exception:
            pass

    # Everything below lives in a single full-width, 1-column table whose only
    # visible divisions are Word's faint gridlines (the boxed look) — exactly
    # like the reference. Each logical block (intro, date, section heading,
    # article) is its own row.
    table = doc.add_table(rows=1, cols=1)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    _no_table_borders(table)
    table.autofit = True

    _first_row = [True]

    def new_block():
        if _first_row[0]:
            _first_row[0] = False
            return _Block(table.rows[0].cells[0])
        return _Block(table.add_row().cells[0])

    # 2) Intro line, with the submit email as a mailto link.
    if intro is None:
        intro = (
            f"The {brand} News Monitor is a snapshot of stories about the company "
            "and its brands; its industrial and peer benchmark group; and its "
            "relevant trends and issues. If you have a story for the daily "
            "monitor, please submit it to "
        )
    blk = new_block()
    ip = blk.add_paragraph()
    _tight(ip, space_before=0, space_after=4)
    _add_run(ip, intro)
    if submit_email:
        _add_external_hyperlink(ip, f"mailto:{submit_email}", submit_email)
        _add_run(ip, ".")

    # 3) Report date — subtle gray small caps.
    blk = new_block()
    dp = blk.add_paragraph()
    _tight(dp, space_before=2, space_after=2)
    _add_run(dp, _fmt_report_date(report_date), color=DATE_COLOR, small_caps=True)

    # Spacer row between the date and the first section (matches the reference).
    _blank_line(new_block())

    # 4) Sections.
    for section in sections:
        articles = section.get("articles", [])

        blk = new_block()
        head = blk.add_paragraph()
        _tight(head, space_before=4, space_after=2)
        # Bold, all-caps, purple, 12pt — mirrors the reference "Heading 1".
        _add_run(head, str(section["name"]), bold=True, caps=True,
                 color=HEADING_COLOR, size=HEADING_SZ)

        if not articles:
            empty = blk.add_paragraph()
            _tight(empty)
            _add_run(empty, "No articles in this section.", italic=True)

        for art in articles:
            blk = new_block()

            # Source label — bold + italic, uppercase.
            sp = blk.add_paragraph()
            _tight(sp, space_before=4, space_after=1)
            _add_run(sp, _article_source(art), bold=True, italic=True, color=BLACK)

            # Title — clickable when a URL is present, else plain text.
            title = title_case(art.get("title") or "Untitled")
            url = art.get("url") or ""
            tp = blk.add_paragraph()
            _tight(tp, space_before=0, space_after=8)
            if url:
                _add_external_hyperlink(tp, url, title)
            else:
                _add_run(tp, title)

            # "Summary:" bold label on its own line, then the summary body.
            lp = blk.add_paragraph()
            _tight(lp, space_before=0, space_after=8)
            _add_run(lp, "Summary: ", bold=True)

            content = art.get("summary") or art.get("content") or ""
            if content:
                cp = blk.add_paragraph()
                _tight(cp, space_before=0, space_after=8)
                _add_run(cp, str(content).strip())

            # "This story was also covered by ..." — comma-separated links
            # (black, bold, underlined — color="auto" over the Hyperlink style).
            similar = art.get("similar_articles") or {}
            similar = similar if isinstance(similar, dict) else {}
            if similar:
                ap = blk.add_paragraph()
                _tight(ap, space_before=0, space_after=8)
                _add_run(ap, "This story was also covered by ")
                items = list(similar.items())
                for j, (domain, durl) in enumerate(items):
                    if j:
                        _add_run(ap, ", ")
                    if durl:
                        _add_external_hyperlink(ap, durl, domain, color=BLACK)
                    else:
                        _add_run(ap, domain, bold=True)
                _add_run(ap, ".")

            # "Translated from <language>." for foreign-language items.
            lang = _translated_from(art)
            if lang:
                tlp = blk.add_paragraph()
                _tight(tlp, space_before=0, space_after=2)
                _add_run(tlp, f"Translated from {lang}.", bold=True, italic=True, color=BLACK)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
