"""Build an Otsuka-style media-coverage report (.docx) from dashboard data.

The output mirrors the reference report (`dumps/otsuka_report_example.docx`): an
Arial 10pt "No Spacing" document with a logo banner, laid out as:

    Friday, June 5, 2026                                    (right-aligned, bold)

    Otsuka Mentions | Treatments and Pipeline Mentions | NSA | ...   (index links)

    ┌───────────────────────────────────────────────────────────┐
    │ Otsuka Mentions                          (white text, blue bar) │
    └───────────────────────────────────────────────────────────┘
    Otsuka Says Kidney Disease Drug Preserves Function ...   (blue title link)
    Reuters, June 4, 2026                                    (bold)
    By Siddhi Mahatole and Sriparna Roy                      (bold)

    The article reports that ... The article mentions Otsuka but there was no
    mention of Rexulti and Lundbeck in this coverage.        (summary paragraph)

    Similar coverage in Endpoints News, STAT, and ClinicalTrials Arena.  (italic links)

                              [Back to Top]                  (centered, links up)

Each section heading is a shaded blue bar with a bookmark; the index at the top
links down to it and every section ends with a "[Back to Top]" link back up.
Sections with no articles render a single "N/A" line. See
`build_media_monitoring_report` for the entry point. `sections_from_charts_data`
(shared with the BeOne report) extracts the ordered sections from charts data.
"""

from __future__ import annotations

import io
import os
from datetime import date, datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches, Emu

# Re-export so callers can build sections the same way as the other reports.
from reports_helpers.beone_report import (  # noqa: F401
    sections_from_charts_data, title_case, clean_author,
)

# --- constants ---------------------------------------------------------------

FONT = "Arial"
FONT_SZ = Pt(10)  # matches the reference document's "No Spacing" body text.
FOOTER_SZ = Pt(9)  # the reference footnotes are 9pt.

LOGO_PATH = os.path.join(os.path.dirname(__file__), "assets", "otsuka_logo.png")
# Reference logo dimensions (EMU) — ~1.75in x 0.66in.
LOGO_WIDTH = Emu(1685925)
LOGO_HEIGHT = Emu(638175)

# Hyperlinks use Word's built-in "Hyperlink" look (blue, underlined).
LINK_BLUE = RGBColor(0x05, 0x63, 0xC1)
# Section heading bars: white text on a light-blue fill.
HEADING_FILL = "51A7F9"
HEADING_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
# The "[Back to Top]" link is a near-black grey in the reference.
BACKTOTOP_GREY = RGBColor(0x24, 0x24, 0x24)

TOP_BOOKMARK = "Backtotop"

FOOTER_COPYRIGHT = (
    "Materials included or referenced in this report may be the subject of "
    "copyright. Users must verify that they have sufficient permissions, rights, "
    "and/or licenses prior to any copying or redistribution. Please consult your "
    "copyright compliance guidelines and personnel with questions."
)
FOOTER_TIP = "Tip: Having trouble viewing a webpage? Try using a different browser."


# --- low-level docx helpers --------------------------------------------------

def _style_run(run, *, bold=False, italic=False, color=None, size=FONT_SZ,
               underline=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.font.underline = underline
    if color is not None:
        run.font.color.rgb = color
    return run


def _add_run(paragraph, text, **kw):
    return _style_run(paragraph.add_run(text), **kw)


def _add_external_hyperlink(paragraph, url, text, *, bold=False, italic=False,
                            color=LINK_BLUE, underline=True):
    """Add a clickable external hyperlink run to `paragraph`."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = paragraph.add_run(text)
    _style_run(run, bold=bold, italic=italic, color=color, underline=underline)
    hyperlink.append(run._element)
    paragraph._p.append(hyperlink)
    return hyperlink


def _add_internal_hyperlink(paragraph, anchor, text, *, bold=False, italic=False,
                            color=None, underline=False):
    """Add an in-document hyperlink run that jumps to a bookmark `anchor`."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = paragraph.add_run(text)
    _style_run(run, bold=bold, italic=italic, color=color, underline=underline)
    hyperlink.append(run._element)
    paragraph._p.append(hyperlink)
    return hyperlink


def _add_bookmark(paragraph, name, bm_id):
    """Wrap the start of `paragraph` in a Word bookmark named `name`."""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bm_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bm_id))
    # Insert the start before any runs and the end right after it.
    paragraph._p.insert(0, start)
    start.addnext(end)


def _shade_paragraph(paragraph, fill):
    """Give `paragraph` a solid background fill (the section-heading bar)."""
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def _tight(paragraph, *, space_before=0, space_after=0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return paragraph


def _blank_line(doc):
    p = doc.add_paragraph()
    _tight(p)
    _add_run(p, "")
    return p


# --- value formatting --------------------------------------------------------

def _parse_date(value):
    if value is None or value == "":
        return None
    if isinstance(value, (datetime, date)):
        return value
    s = str(value).strip()
    # Normalize an ISO 8601 'Z' / timezone suffix so fromisoformat accepts it.
    iso = s.replace("Z", "+00:00")
    try:
        return datetime.fromisoformat(iso)
    except ValueError:
        pass
    for fmt in ("%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d",
                "%Y/%m/%d", "%d-%m-%Y", "%m/%d/%Y",
                "%d-%m-%y", "%m-%d-%y",
                "%m.%d.%Y", "%m.%d.%y", "%d.%m.%Y", "%Y.%m.%d"):
        try:
            return datetime.strptime(s[:len(fmt) + 2].strip(), fmt)
        except ValueError:
            continue
    try:
        return datetime.strptime(s[:10], "%Y-%m-%d")
    except ValueError:
        return None


def _fmt_long_date(value):
    """Format an article date as the reference's `Month D, YYYY` (e.g. June 4, 2026)."""
    dt = _parse_date(value)
    # `%-d` is not portable (fails on Windows), so build the day without padding.
    return f"{dt.strftime('%B')} {dt.day}, {dt.year}" if dt else ""


def _fmt_weekday_date(value):
    """Format the header date as `Weekday, Month D, YYYY` (e.g. Friday, June 5, 2026)."""
    dt = _parse_date(value)
    return f"{dt.strftime('%A, %B')} {dt.day}, {dt.year}" if dt else ""


def _article_source(article):
    src = (
        article.get("domain")
        or article.get("source")
        or article.get("publication")
        or article.get("domain_name")
        or "Unknown source"
    )
    return str(src).strip()


def _bookmark_name(index):
    """A valid, unique Word bookmark name for the section at `index`."""
    return f"section_{index}"


def _summary_text(article):
    """Flatten an article's summary points into one flowing paragraph.

    The synthesizer stores the point-wise summary as newline-joined text in
    ``content``; the reference renders it as a single paragraph, so we join the
    points with spaces and append the brand-mention note as the closing sentence.
    """
    content = article.get("summary") or article.get("content") or ""
    points = [p.strip() for p in str(content).replace("\r\n", "\n").split("\n")]
    body = " ".join(p for p in points if p)

    note = str(article.get("mention_note") or "").strip()
    if note and note not in body:
        body = f"{body} {note}".strip() if body else note
    return body


def _join_similar(paragraph, similar):
    """Render 'Similar coverage in A, B, and C.' with each source hyperlinked."""
    items = [(str(k).strip(), v) for k, v in similar.items() if str(k).strip()]
    if not items:
        return
    _add_run(paragraph, "Similar coverage in ", italic=True)
    last = len(items) - 1
    for j, (domain, durl) in enumerate(items):
        if j:
            if len(items) > 2:
                _add_run(paragraph, ", and " if j == last else ", ", italic=True)
            else:
                _add_run(paragraph, " and ", italic=True)
        if durl:
            _add_external_hyperlink(paragraph, str(durl), domain, italic=True)
        else:
            _add_run(paragraph, domain, italic=True)
    _add_run(paragraph, ".", italic=True)


def _back_to_top(doc):
    """Centered "[Back to Top]" link that jumps to the top-of-document bookmark."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight(p, space_after=2)
    _add_run(p, "[", color=BACKTOTOP_GREY)
    _add_internal_hyperlink(p, TOP_BOOKMARK, "Back to Top", color=BACKTOTOP_GREY, underline=True)
    _add_run(p, "]", color=BACKTOTOP_GREY)
    _blank_line(doc)
    return p


# --- report builder ----------------------------------------------------------

def build_media_monitoring_report(
    sections,
    *,
    report_date=None,
    brand="Otsuka",
    title=None,
    intro=None,
    summary_paragraphs=None,
    signoff=None,
    brand_keywords=None,
    competitor_keywords=None,
):
    """Render the Otsuka news-coverage report and return the .docx as bytes.

    Args:
        sections: list of ``{"name": str, "articles": [article dict, ...]}`` in
            display order. Each article dict may contain title, content/summary,
            mention_note, url, domain/source, author, date and similar_articles
            ({domain: url}).
        report_date: date shown in the header line. Defaults to the latest article
            date, then today.
        brand: brand name (kept for signature compatibility; unused in this layout).
        title / intro: unused; kept for signature compatibility with the other reports.
        summary_paragraphs: optional executive-summary paragraphs rendered between
            the index and the first section; omit for the plain reference layout.
        signoff / brand_keywords / competitor_keywords: unused; kept for
            signature compatibility with the other report builders.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for style_name in ("Normal", "No Spacing"):
        try:
            style = doc.styles[style_name]
            style.font.name = FONT
            style.font.size = FONT_SZ
        except KeyError:
            pass

    sections = list(sections or [])

    if report_date is None:
        latest = None
        for section in sections:
            for art in section.get("articles", []):
                dt = _parse_date(art.get("date"))
                if dt and (latest is None or dt > latest):
                    latest = dt
        report_date = latest or datetime.today()

    bm_ids = iter(range(1, 100_000))

    # 1) Logo banner.
    if os.path.exists(LOGO_PATH):
        lp = doc.add_paragraph()
        _tight(lp, space_after=4)
        run = lp.add_run()
        try:
            run.add_picture(LOGO_PATH, width=LOGO_WIDTH, height=LOGO_HEIGHT)
        except Exception:
            pass

    # 2) Header date — right-aligned, bold.
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _tight(dp)
    _add_run(dp, _fmt_weekday_date(report_date), bold=True)

    # 3) "Back to Top" anchor lives on the blank line under the date.
    top_anchor = _blank_line(doc)
    _add_bookmark(top_anchor, TOP_BOOKMARK, next(bm_ids))

    # 4) Index — every section name as a blue, underlined in-document link,
    #    separated by " | ".
    nav = doc.add_paragraph()
    _tight(nav)
    for i, section in enumerate(sections):
        if i:
            _add_run(nav, " | ", bold=True)
        _add_internal_hyperlink(
            nav, _bookmark_name(i), str(section.get("name", "")),
            bold=True, color=LINK_BLUE, underline=True,
        )
    _blank_line(doc)

    # 5) Optional executive-summary paragraphs (omitted in the plain reference).
    for para in (summary_paragraphs or []):
        if not str(para).strip():
            continue
        sp = doc.add_paragraph()
        _tight(sp)
        _add_run(sp, str(para).strip())
        _blank_line(doc)

    # 6) Sections.
    for i, section in enumerate(sections):
        articles = section.get("articles", [])

        head = doc.add_paragraph()
        _tight(head, space_before=2, space_after=2)
        _shade_paragraph(head, HEADING_FILL)
        _add_bookmark(head, _bookmark_name(i), next(bm_ids))
        _add_run(head, str(section.get("name", "")), bold=True, color=HEADING_TEXT)
        _blank_line(doc)

        if not articles:
            na = doc.add_paragraph()
            _tight(na, space_after=2)
            _add_run(na, "N/A")
            _blank_line(doc)
            _back_to_top(doc)
            continue

        for art in articles:
            # Title line — bold blue hyperlink to the article (plain if no url).
            tp = doc.add_paragraph()
            _tight(tp)
            title_text = title_case(art.get("title") or "Untitled")
            url = art.get("url") or ""
            if url:
                _add_external_hyperlink(tp, str(url), title_text, bold=True)
            else:
                _add_run(tp, title_text, bold=True)

            # Source + date — bold, e.g. "Reuters, June 4, 2026".
            source = _article_source(art)
            long_date = _fmt_long_date(art.get("date"))
            sd = doc.add_paragraph()
            _tight(sd)
            _add_run(sd, f"{source}, {long_date}" if long_date else source, bold=True)

            # Author — bold, e.g. "By Siddhi Mahatole and Sriparna Roy".
            author = clean_author(art.get("author"))
            ap = doc.add_paragraph()
            _tight(ap)
            _add_run(ap, f"By {author}" if author else "N/A", bold=True)

            _blank_line(doc)

            # Summary paragraph (points flattened; mention note as closing sentence).
            body = _summary_text(art)
            if body:
                bp = doc.add_paragraph()
                _tight(bp)
                _add_run(bp, body)
                _blank_line(doc)

            # "Similar coverage in ..." — italic, comma-separated links.
            similar = art.get("similar_articles") or {}
            if isinstance(similar, dict) and similar:
                simp = doc.add_paragraph()
                _tight(simp)
                _join_similar(simp, similar)
                _blank_line(doc)

        _back_to_top(doc)

    # 7) Footer notes.
    cp = doc.add_paragraph()
    _tight(cp, space_before=10)
    _add_run(cp, FOOTER_COPYRIGHT, italic=True, size=FOOTER_SZ)
    tp = doc.add_paragraph()
    _tight(tp)
    _add_run(tp, FOOTER_TIP, italic=True, size=FOOTER_SZ)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
