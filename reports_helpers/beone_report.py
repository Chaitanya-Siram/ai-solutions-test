"""Build a BeOne-style media-monitoring report (.docx) from dashboard data.

The output mirrors the reference report (`beone_report_example.docx`): an Arial
10pt document with a logo banner, a report date, an intro line, a section
navigation strip, then one shaded block per media-monitoring section followed by
its articles (clickable title, italic "date, source" line, summary) and a
"[Back to Top]" link. See `build_media_monitoring_report` for the entry point.
"""

from __future__ import annotations

import io
import os
import re
from datetime import date, datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Emu, Inches
from agents.tagging_agent.tagging_common import NO_SECTION

# --- constants ---------------------------------------------------------------

FONT = "Arial"
FONT_SZ = Pt(10)  # 20 half-points, matching the reference document.

# Section header shading: the first section is highlighted red, the rest blue
# (exactly as the reference report distinguishes "New & Noteworthy News").
FIRST_SECTION_FILL = "C00000"
SECTION_FILL = "365F91"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LINK_BLUE = RGBColor(0x05, 0x63, 0xC1)

TOP_BOOKMARK = "report_top"

LOGO_PATH = os.path.join(os.path.dirname(__file__), "assets", "beone_logo.jpeg")
# Reference banner dimensions (EMU) — ~6.27in x 1.26in.
LOGO_WIDTH = Emu(5731510)
LOGO_HEIGHT = Emu(1155065)

FOOTER_COPYRIGHT = (
    "Materials included or referenced in this report may be the subject of "
    "copyright. Users must verify that they have sufficient permissions, rights, "
    "and/or licenses prior to any copying or redistribution. Please consult your "
    "copyright compliance guidelines and personnel with questions."
)
FOOTER_TIP = "Tip: Having trouble viewing a webpage? Try using a different browser."


# --- shared helpers ----------------------------------------------------------

def title_case(text):
    """Capitalize the first letter of every word in an article headline.

    Not `str.title()` or `string.capwords()`, which lowercase the rest of each
    word and so would turn "FDA" into "Fda" and "340B" into "340b" ( `.title()`
    also capitalizes after apostrophes, giving "Otsuka'S"). Words that already
    contain an uppercase letter (acronyms, brand names) are left as written.

    Args:
        text: The headline.

    Returns:
        The headline with each word capitalized.
    """
    words = []
    for word in str(text).split(" "):
        if word and word.islower():
            word = word[0].upper() + word[1:]
        words.append(word)
    return " ".join(words)


# Byline junk newspaper3k picks up from nav links, promos and author bios.
# Matched against a whole comma-separated part, case-insensitively.
_AUTHOR_JUNK_WORDS = (
    "sign up", "subscribe", "newsletter", "log in", "login", "sign in",
    "share", "follow", "read more", "click here", "contact us", "about us",
    "privacy", "terms", "cookie", "advertise", "correspondent", "staff",
    "editor", "editorial", "newsroom", "press release", "admin",
    "authority", "reporter", "contributor", "guest", "opinion", "columnist",
)

# A byline part that reads like a bio sentence ("Chris Jacobs Is Founder"),
# a job title, or an organization rather than a person's name.
_AUTHOR_BIO_WORDS = (
    " is ", " was ", " has ", " have ", " who ", " of the ", " at the ",
    "founder", "ceo", "cto", "cfo", "president", "director", "author of",
    "professor", "chairman", "partner at", "research group", "institute",
    "university", "llc", "inc.", "ltd", "group", "company", "http",
)

_MAX_AUTHOR_WORDS = 5


def _is_person_name(part):
    """Whether one comma-separated byline part looks like a person's name."""
    low = f" {part.lower()} "
    if any(w in low for w in _AUTHOR_JUNK_WORDS):
        return False
    if any(w in low for w in _AUTHOR_BIO_WORDS):
        return False
    words = part.split()
    # A name is a couple of words; anything longer is a title or a sentence.
    if not (1 < len(words) <= _MAX_AUTHOR_WORDS):
        return False
    return all(w[0].isalpha() for w in words if w)


def clean_author(value):
    """Clean a scraped byline down to the real person names it contains.

    newspaper3k harvests bylines heuristically and often returns nav links
    ("Sign Up Now", "Employment Authority") or the author's bio sentence
    ("Chris Jacobs Is Founder, Ceo Of Juniper Research Group"). Each
    comma-separated part is kept only when it still looks like a person's name.

    Args:
        value: The raw author string.

    Returns:
        The cleaned byline, or "" when nothing survives.
    """
    text = str(value or "").strip()
    if text.lower().startswith("by "):
        text = text[3:]
    # Bylines join names with either commas or "and" ("Jane Doe and John Roe").
    parts = [p.strip(" \t\r\n-|·•") for p in re.split(r",|\band\b", text)]
    seen, names = set(), []
    for part in parts:
        if not _is_person_name(part):
            continue
        low = part.lower()
        if low not in seen:
            seen.add(low)
            names.append(part)
    if len(names) > 1:
        return f"{', '.join(names[:-1])} and {names[-1]}"
    return names[0] if names else ""


# --- low-level docx helpers --------------------------------------------------

def _style_run(run, *, bold=False, italic=False, color=None, size=FONT_SZ):
    run.font.name = FONT
    # Ensure the font applies to complex/east-asian scripts too.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return run


def _add_run(paragraph, text, **kw):
    return _style_run(paragraph.add_run(text), **kw)


def _shade_paragraph(paragraph, fill):
    """Apply a solid background fill to a paragraph (full-width shaded block)."""
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def _add_bookmark(paragraph, name, bookmark_id):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_external_hyperlink(paragraph, url, text, *, bold=False, color=LINK_BLUE,
                            underline=True):
    """Add a clickable external hyperlink run to `paragraph`."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = paragraph.add_run(text)
    run.font.underline = underline
    _style_run(run, bold=bold, color=color)
    hyperlink.append(run._element)
    paragraph._p.append(hyperlink)
    return hyperlink


def _add_internal_hyperlink(paragraph, anchor, text, *, bold=False, color=LINK_BLUE):
    """Add a clickable link that jumps to a bookmark within the document."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = paragraph.add_run(text)
    _style_run(run, bold=bold, color=color)
    hyperlink.append(run._element)
    paragraph._p.append(hyperlink)
    return hyperlink


def _tight(paragraph, *, space_before=2, space_after=2):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return paragraph


def _blank_line(doc):
    """Add an empty paragraph — a single blank line at the document font size."""
    p = doc.add_paragraph()
    _tight(p, space_before=0, space_after=0)
    _add_run(p, "")
    return p


# --- value formatting --------------------------------------------------------

def _slug(name):
    keep = "".join(c if c.isalnum() else "_" for c in str(name))
    return f"sec_{keep.strip('_').lower()}"[:40]


def _parse_date(value):
    if value is None or value == "":
        return None
    if isinstance(value, (datetime, date)):
        return value
    s = str(value)
    for fmt in ("%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d",
                "%Y/%m/%d", "%d-%m-%Y", "%m/%d/%Y"):
        try:
            return datetime.strptime(s[:len(fmt) + 2].strip(), fmt)
        except ValueError:
            continue
    # Fall back to the leading YYYY-MM-DD if present.
    try:
        return datetime.strptime(s[:10], "%Y-%m-%d")
    except ValueError:
        return None


def _fmt_source_date(value):
    """Format an article date as the reference's `YYYY/MM/DD` source-line prefix."""
    dt = _parse_date(value)
    return dt.strftime("%Y/%m/%d") if dt else ""


def _fmt_report_date(value):
    """Format the report header date, e.g. `Thursday, May 14, 2026`."""
    dt = _parse_date(value)
    if not dt:
        return ""
    # Cross-platform: avoid %-d / %#d differences.
    return f"{dt.strftime('%A, %B')} {dt.day}, {dt.year}"


def _article_source(article):
    return (
        article.get("domain")
        or article.get("source")
        or article.get("publication")
        or "Unknown source"
    )


# --- report builder ----------------------------------------------------------

def build_media_monitoring_report(
    sections,
    *,
    report_date=None,
    brand="BeOne",
    intro=None,
    include_logo=True,
):
    """Render the report and return the .docx as bytes.

    Args:
        sections: list of ``{"name": str, "articles": [article dict, ...]}`` in
            display order. Each article dict may contain title, content/summary,
            url, domain/source, date, reach, sentiment, priority and
            similar_articles ({domain: url}).
        report_date: date shown in the header (date/datetime/str). Defaults to the
            most recent article date, then today.
        brand: brand name used in the intro line.
        intro: override the default intro sentence.
        include_logo: embed the banner logo if the asset is present.
    """
    doc = Document()

    # Use Word's "Normal" margin preset (1" on every side) rather than the
    # python-docx default (the old Office 1" top/bottom, 1.25" left/right).
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document-wide default font.
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = FONT_SZ

    # Derive the report date from the latest article when not supplied.
    if report_date is None:
        latest = None
        for section in sections:
            for art in section.get("articles", []):
                dt = _parse_date(art.get("date"))
                if dt and (latest is None or dt > latest):
                    latest = dt
        report_date = latest or datetime.today()

    bookmark_id = 0

    # 1) Logo banner.
    if include_logo and os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        _tight(p, space_before=0, space_after=4)
        run = p.add_run()
        try:
            run.add_picture(LOGO_PATH, width=LOGO_WIDTH, height=LOGO_HEIGHT)
        except Exception:
            pass

    # 2) Report date — right-aligned, bold.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _tight(p)
    _add_run(p, _fmt_report_date(report_date), bold=True)

    # 3) Intro line.
    if intro is None:
        intro = (
            f"Below, please find {brand} announcements and news, and a sampling "
            "of industry news coverage."
        )
    p = doc.add_paragraph()
    _tight(p)
    _add_run(p, intro)

    # Blank line after the intro line.
    _blank_line(doc)

    # 4) Section navigation strip (each name links to its section) + top anchor.
    nav = doc.add_paragraph()
    _tight(nav)
    bookmark_id += 1
    _add_bookmark(nav, TOP_BOOKMARK, bookmark_id)
    for i, section in enumerate(sections):
        if i:
            _add_run(nav, " | ", bold=True)
        _add_internal_hyperlink(nav, _slug(section["name"]), section["name"], bold=True)

    # 5) Sections.
    for i, section in enumerate(sections):
        articles = section.get("articles", [])
        fill = FIRST_SECTION_FILL if i == 0 else SECTION_FILL

        head = doc.add_paragraph()
        _tight(head, space_before=8, space_after=2)
        _shade_paragraph(head, fill)
        bookmark_id += 1
        _add_bookmark(head, _slug(section["name"]), bookmark_id)
        _add_run(head, section["name"], bold=True, color=WHITE)

        if not articles:
            empty = doc.add_paragraph()
            _tight(empty)
            _add_run(empty, "No articles in this section.", italic=True)

        for art in articles:
            title = title_case(art.get("title") or "Untitled")
            url = art.get("url") or ""

            # Title — bold, clickable when a URL is present.
            tp = doc.add_paragraph()
            _tight(tp, space_before=6, space_after=1)
            if url:
                _add_external_hyperlink(tp, url, title, bold=True, color=LINK_BLUE)
            else:
                _add_run(tp, title, bold=True)

            # Source line — "YYYY/MM/DD, Source", italic.
            src_date = _fmt_source_date(art.get("date"))
            source = _article_source(art)
            meta = ", ".join(x for x in (src_date, source) if x)
            if meta:
                mp = doc.add_paragraph()
                _tight(mp, space_before=0, space_after=1)
                _add_run(mp, meta, italic=True)

            # Prefer the AI-generated summary; fall back to full content when absent.
            # Similar coverage (if any) is appended to this same paragraph as a
            # trailing sentence — not a separate paragraph.
            content = art.get("summary") or art.get("content") or ""
            similar = art.get("similar_articles") or {}
            similar = similar if isinstance(similar, dict) else {}

            if content or similar:
                cp = doc.add_paragraph()
                _tight(cp, space_before=0, space_after=2)
                if content:
                    _add_run(cp, str(content).strip())

                if similar:
                    items = list(similar.items())
                    last = len(items) - 1
                    if content:
                        _add_run(cp, " ")
                    _add_run(cp, "Similar News was reported by ")
                    for j, (domain, durl) in enumerate(items):
                        if j:
                            # Comma-separate, with "and" before the final link
                            # (Oxford comma when there are three or more).
                            if len(items) > 2:
                                _add_run(cp, ", and " if j == last else ", ")
                            else:
                                _add_run(cp, " and ")
                        if durl:
                            _add_external_hyperlink(cp, durl, domain, color=LINK_BLUE)
                        else:
                            _add_run(cp, domain)
                    _add_run(cp, ".")

            # Blank line after each article.
            _blank_line(doc)

        # [Back to Top] — centered link to the navigation anchor.
        back = doc.add_paragraph()
        back.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _tight(back, space_before=4, space_after=4)
        _add_run(back, " [", bold=True)
        _add_internal_hyperlink(back, TOP_BOOKMARK, "Back to Top", bold=True)
        _add_run(back, "] ", bold=True)

    # 6) Footer notes.
    fp = doc.add_paragraph()
    _tight(fp, space_before=10)
    _add_run(fp, FOOTER_COPYRIGHT)
    tp = doc.add_paragraph()
    _tight(tp)
    _add_run(tp, FOOTER_TIP)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def sections_from_charts_data(charts_data, *, days=None, chart_id="section_articles"):
    """Extract ordered ``[{name, articles}]`` from a charts-data payload.

    `charts_data` is the JSON stored for a session. The media-monitoring section
    table lives under ``charts_data["media_monitoring"]`` as a list of charts; we
    pull the one with ``chart_id == "section_articles"`` whose ``data`` is a
    ``{section_name: [articles]}`` dict. When `days` (a set/list of YYYY-MM-DD
    strings) is given, only articles on those days are kept.
    """
    arr = charts_data.get("media_monitoring") if isinstance(charts_data, dict) else None
    chart = None
    if isinstance(arr, list):
        chart = next((c for c in arr if isinstance(c, dict) and c.get("chart_id") == chart_id), None)
    data = chart.get("data") if isinstance(chart, dict) else None
    if not isinstance(data, dict):
        return []

    day_filter = set(days) if days else None
    sections = []
    for name, articles in data.items():
        if str(name).strip().lower() == NO_SECTION.lower():
            continue
        if not isinstance(articles, list):
            articles = []
        if day_filter is not None:
            articles = [
                a for a in articles
                if str(a.get("date") or "")[:10] in day_filter
            ]
        sections.append({"name": name, "articles": articles})
    return sections
